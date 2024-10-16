from flask import Flask, render_template, request, send_file, jsonify
import json
from docx import Document
from io import BytesIO
# Memory Check
# import psutil
import os

app = Flask(__name__)

# def log_memory_usage(label):
#     process = psutil.Process(os.getpid())
#     memory_info = process.memory_info()
#     print(f"{label} - Memory Usage: {memory_info.rss / (1024 * 1024)} MB")  # Memory usage in MB

# Increase in memory usage indicates that BytesIO is working as expected.
# Possibly need to integrate GC (Garbage Collection) manually to ensure dumping of information.

# Import JSON Data
JSON_FILE = "booking_data/BookingData.json"

def Load_JSON():
    try:
        with open(JSON_FILE, "r") as file:
            appData = json.load(file)
            print("SUCCESS: JSON Data Loaded")
            return appData
    except Exception as e:
        print(f"ERROR: Unable To Load JSON Data: {e}")
        return None

appData = Load_JSON()

@app.route('/')
def index():
    return render_template('index.html', data=appData)

# API to provide mapping based on party type
@app.route('/room_mapping', methods=['POST'])
def room_mapping():
    party_type = request.json.get('party_type')
    party_room = request.json.get('party_room')

    # Check if a party type is provided
    if party_type:
        activity_rooms = appData['ACTIVITY_ROOM_MAPPING'].get(party_type, [])
        return jsonify(activity_rooms)

    # If a party room is provided
    if party_room:
        food_rooms = appData['FOOD_ROOM_MAPPING'].get(party_room, [])
        return jsonify(food_rooms)

    return jsonify([])

@app.route('/generate_document', methods=['POST'])
def generate_document():
    # log_memory_usage("Pre Generation")
    # Extract form data
    customer_name = request.form['customer_name']
    customer_email = request.form['customer_email']
    customer_phone = request.form['customer_phone']
    child_name = request.form['child_name']
    child_age = request.form['child_age']
    party_date = request.form['party_date']
    party_start_time = request.form['party_start_time']
    party_end_time = request.form['party_end_time']
    date_booked = request.form['date_booked']
    staff_member = request.form['staff_member']

    party_type = request.form['party_type']
    party_room = request.form['party_room']
    party_food_room = request.form['party_food_room']

    # Split party type and cost
    party_activity, party_cost = party_type.split(": £")

    # Document: Web Information
    CUSTOMER_INFORMATION = {
        "CUSTOMER_NAME": customer_name,
        "CUSTOMER_EMAIL": customer_email,
        "CUSTOMER_NUMBER": customer_phone
    }
    CHILD_INFORMATION = {
        "CHILD_NAME": child_name,
        "CHILD_AGE": child_age
    }
    PARTY_INFORMATION = {
        "PARTY_DATE": party_date,
        "PARTY_START_TIME": party_start_time,
        "PARTY_END_TIME": party_end_time,
        "PARTY_TYPE": party_activity,
        "PARTY_COST": party_cost,
        "PARTY_ROOM": party_room,
        "PARTY_FOOD_ROOM": party_food_room,
        "MAX_CHILDREN": appData["MAXIMUM_CHILDREN"][party_activity]
    }
    ADMIN_INFORMATION = {
        "CUSTOMER_FIRST_NAME": customer_name.split(" ")[0],
        "DATE_BOOKED": date_booked,
        "STAFF_MEMBER": staff_member
    }

    # Load Template
    doc = Document(appData["TEMPLATE_DOCUMENT"])

    # Replace Keywords
    for paragraph in doc.paragraphs:
        for key, value in {**CUSTOMER_INFORMATION, **CHILD_INFORMATION, **PARTY_INFORMATION, **ADMIN_INFORMATION}.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text
                for key, value in {**CUSTOMER_INFORMATION, **CHILD_INFORMATION, **PARTY_INFORMATION, **ADMIN_INFORMATION}.items():
                    if key in cell_text:
                        cell_text = cell_text.replace(key, str(value))
                cell.text = cell_text

    # Save The Document - BytesIO uses memory as a temporary storage.
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    download_filename = f"{CUSTOMER_INFORMATION['CUSTOMER_NAME']} - {party_activity} - Party Confirmation.docx"

    # log_memory_usage("Post Generation")

    # Send Document to user for download.
    return send_file(
        doc_io,
        as_attachment=True,
        download_name=download_filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )

if __name__ == "__main__":
    app.run(debug=True)
